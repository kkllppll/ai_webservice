from flask import Blueprint, jsonify, request, g
from app.services.sentiment_service import analyze_sentiment
from app.services.topic_service import classify_topic
from app.services.toxicity_service import detect_toxicity
from app.services.openai_service import generate_summary  
from app.utils.auth_decorator import auth_required
from app.services.firebase_service import get_analyses
from app.services.firebase_service import delete_user_analyses 
from app.services.firebase_service import delete_user_analysis_by_id
import os
from openai import OpenAI
import fitz  
import json
import tempfile
import os
import docx
import whisper
import pandas as pd



bp = Blueprint('api', __name__)

@bp.route('/ping')
def ping():
    return jsonify({"status": "ok"})

@bp.route('/analyze/sentiment', methods=['POST'])
def sentiment():
    data = request.get_json()
    text = data.get("text", "")

    if not text.strip():
        return jsonify({"error": "Text is required"}), 400

    try:
        result = analyze_sentiment(text)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    

@bp.route('/analyze/topic', methods=['POST'])
def analyze_topic():
    data = request.get_json() or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "Empty text"}), 400

    result = classify_topic(text)
    return jsonify({"topics": result})



@bp.route('/analyze/toxicity', methods=['POST'])
def analyze_toxicity_endpoint():
    data = request.get_json() or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "Empty text"}), 400

    result = detect_toxicity(text)
    return jsonify({"toxicity": result})


@bp.route('/analyze/summary', methods=['POST'])
def get_summary():
    data = request.get_json() or {}
    text = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "Empty text"}), 400

    result = generate_summary(text)
    return jsonify({"summary": result})



@bp.route('/auth/protected', methods=['GET'])
@auth_required
def protected_route():
    return jsonify({
        "message": "Access granted",
        "user_id": g.user.get("uid"),
        "email": g.user.get("email")
    })


@bp.route('/analyses', methods=['GET'])
@auth_required
def get_analyses_route():
    user_id = g.user.get("uid")

    try:
        analyses = get_analyses(user_id)
        return jsonify(analyses), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@bp.route('/delete_history', methods=['DELETE'])
@auth_required
def delete_history():
    user_id = g.user.get("uid")
    try:
        delete_user_analyses(user_id)
        return jsonify({"message": "History deleted successfully"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@bp.route('/analyze_base', methods=['POST'])
def analyze_base():
    data = request.get_json()
    text = data.get("text", "").strip()

    if not text:
        return jsonify({"error": "Text is required"}), 400

    try:
        sentiment_result = analyze_sentiment(text)
        toxicity_result = detect_toxicity(text)

        return jsonify({
            "results": {
                "sentiment": sentiment_result,
                "toxicity": toxicity_result
            }
        }), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@bp.route('/delete_analysis/<analysis_id>', methods=['DELETE'])
@auth_required
def delete_single_analysis(analysis_id):
    user_id = g.user.get("uid")

    try:
        success = delete_user_analysis_by_id(user_id, analysis_id)

        if success:
            return jsonify({"message": "Аналіз видалено"}), 200
        else:
            return jsonify({"error": "Аналіз не знайдено або не належить користувачу"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@bp.route('/regenerate_text', methods=['POST'])
def regenerate_text():
    data = request.get_json()
    prompt = data.get("prompt", "")

    if not prompt.strip():
        return jsonify({"error": "Prompt is empty"}), 400

    try:
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))  

        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.9,
            max_tokens=512
        )

        generated = response.choices[0].message.content
        return jsonify({"text": generated})
    except Exception as e:
        print("OpenAI error:", e)
        return jsonify({"error": str(e)}), 500


@bp.route('/analyze_file', methods=['POST'])
@auth_required
def analyze_file():
    file = request.files.get('file')
    if not file:
        return jsonify({"error": "No file provided"}), 400

    filename = file.filename.lower()
    text = ""

    try:
        if filename.endswith('.txt'):
            text = file.read().decode('utf-8')

        elif filename.endswith('.json'):
            obj = json.load(file)
            text = obj.get('text', '') or json.dumps(obj)

        elif filename.endswith('.pdf'):
            doc = fitz.open(stream=file.read(), filetype='pdf')
            text = "\n".join([page.get_text() for page in doc])

        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])

        elif filename.endswith(('.mp3', '.wav', '.m4a')):
            try:
                suffix = os.path.splitext(filename)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    file.save(tmp.name)

                model = whisper.load_model("base")
                result = model.transcribe(tmp.name)
                text = result["text"]

                os.remove(tmp.name)
            except Exception as e:
                print("Whisper error:", e)
                return jsonify({"error": f"Whisper transcription error: {str(e)}"}), 500

        else:
            return jsonify({"error": "Unsupported file format"}), 400

        if not text.strip():
            return jsonify({"error": "Файл не містить тексту"}), 400

        return jsonify({"text": text})

    except Exception as e:
        return jsonify({"error": str(e)}), 500




@bp.route('/analyze_file_full', methods=['POST'])
@auth_required
def analyze_file_full():
    file = request.files.get('file')
    if not file:
        return jsonify({"error": "No file provided"}), 400

    filename = file.filename.lower()
    text = ""

    try:
        if filename.endswith('.txt'):
            text = file.read().decode('utf-8')

        elif filename.endswith('.json'):
            obj = json.load(file)
            text = obj.get('text', '') or json.dumps(obj)

        elif filename.endswith('.csv'):
            df = pd.read_csv(file)
            text = " ".join(map(str, df.values.flatten()))


        elif filename.endswith('.pdf'):
            doc = fitz.open(stream=file.read(), filetype='pdf')
            text = "\\n".join([page.get_text() for page in doc])

        elif filename.endswith('.docx'):
            doc = docx.Document(file)
            text = "\\n".join([p.text for p in doc.paragraphs])

        elif filename.endswith(('.mp3', '.wav', '.m4a')):
            language = request.form.get('language', 'auto')
            suffix = os.path.splitext(filename)[1]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp_path = tmp.name
            tmp.close()

            file.save(tmp_path)

            model = whisper.load_model("base")
            result = model.transcribe(tmp_path, language=language if language != "auto" else None)
            text = result["text"]

            os.remove(tmp_path)


        else:
            return jsonify({"error": "Unsupported file format"}), 400

        if not text.strip():
            return jsonify({"error": "Файл не містить тексту"}), 400

        # АНАЛІЗ
        user_id = g.user.get("uid")
        from app.services.sentiment_service import analyze_sentiment
        from app.services.topic_service import classify_topic
        from app.services.toxicity_service import detect_toxicity
        from app.utils.text_utils import get_preview_text
        from app.services.openai_service import generate_summary
        from app.services.firebase_service import save_analysis

        sentiment = analyze_sentiment(text)
        topics = classify_topic(text)
        toxicity = detect_toxicity(text)
        summary = generate_summary(text)
        preview = get_preview_text(text)

        results = {
            "sentiment": sentiment,
            "topics": topics,
            "toxicity": toxicity
        }

        payload = {
            "preview_text": preview,
            "summary": summary,
            "input_text": text,
            "language": "uk",
            "results": results
        }

        save_analysis(user_id, payload)

        return jsonify(payload), 200

    except Exception as e:
        print("Full file analyze error:", e)
        return jsonify({"error": str(e)}), 500
